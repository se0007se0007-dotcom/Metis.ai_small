"""DOCX 테스트 보고서 생성 — python-docx + matplotlib 차트.

report 데이터(dict)를 받아 점수 차트, 이슈 표, 동적 테스트 표, 스텝별 비용 차트,
토큰 구성 차트, LLM 리뷰, FinOps 텔레메트리 표가 포함된 .docx 를 생성한다.
"""
import io
import os
import time

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x2F, 0x6F, 0xE0)
MUTED = RGBColor(0x6B, 0x76, 0x88)
RED = RGBColor(0xC0, 0x35, 0x2B)
GREEN = RGBColor(0x2E, 0x8B, 0x57)


def _setup_korean_font():
    """matplotlib 한글 폰트: Windows(Malgun Gothic) → Nanum → 기본."""
    for name in ("Malgun Gothic", "NanumGothic", "AppleGothic"):
        if any(f.name == name for f in font_manager.fontManager.ttflist):
            plt.rcParams["font.family"] = name
            plt.rcParams["axes.unicode_minus"] = False
            return True
    return False


HAS_KFONT = _setup_korean_font()


def L(ko: str, en: str) -> str:
    """차트 라벨: 한글 폰트 없으면 영문 폴백(글자 깨짐 방지)."""
    return ko if HAS_KFONT else en


def _kfont(run):
    """run 에 한글(맑은 고딕) East-Asian 폰트를 안전하게 설정 (rPr 없으면 생성)."""
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")


def _set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")


def _h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level > 0 else RGBColor(0x10, 0x18, 0x28)
        _kfont(run)
    return h


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(htext)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def _chart_png(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _score_chart(scores: dict):
    labels = [L("구조", "Structure"), L("안전성", "Safety"), L("테스트", "Tests"), L("문서화", "Docs")]
    vals = [scores["struct"], scores["safety"], scores["test"], scores["doc"]]
    colors = ["#4f8ff7" if v >= 80 else "#e3b341" if v >= 60 else "#f85149" for v in vals]
    fig, ax = plt.subplots(figsize=(6, 2.4))
    bars = ax.barh(labels, vals, color=colors, height=0.55)
    ax.set_xlim(0, 100)
    ax.invert_yaxis()
    ax.bar_label(bars, fmt="%d", padding=4, fontsize=9)
    ax.set_xlabel(L("점수", "Score"), fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=9)
    return _chart_png(fig)


def _cost_chart(steps: list):
    xs = [f"#{s['step']}" for s in steps]
    costs = [s["cost_usd"] for s in steps]
    fig, ax = plt.subplots(figsize=(6, 2.4))
    ax.bar(xs, costs, color="#4f8ff7", width=0.5)
    ax.set_ylabel("USD", fontsize=9)
    ax.set_title(L("스텝별 LLM 비용", "LLM cost per step"), fontsize=10)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=9)
    return _chart_png(fig)


def _token_chart(steps: list):
    cats = [("input_tokens", L("입력", "input"), "#4f8ff7"),
            ("output_tokens", L("출력", "output"), "#9d6ef7"),
            ("cache_read_tokens", L("캐시읽기", "cache_read"), "#3fb97a"),
            ("reasoning_tokens", "reasoning", "#f85149")]
    xs = [f"#{s['step']}" for s in steps]
    fig, ax = plt.subplots(figsize=(6, 2.6))
    bottom = [0] * len(steps)
    for key, label, color in cats:
        vals = [s.get(key, 0) or 0 for s in steps]
        if sum(vals) == 0:
            continue
        ax.bar(xs, vals, bottom=bottom, label=label, color=color, width=0.5)
        bottom = [b + v for b, v in zip(bottom, vals)]
    ax.set_ylabel(L("토큰", "tokens"), fontsize=9)
    ax.set_title(L("스텝별 토큰 구성", "Token composition per step"), fontsize=10)
    ax.legend(fontsize=8, ncols=4, frameon=False)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=9)
    return _chart_png(fig)


def build_docx(data: dict, out_path: str) -> str:
    doc = Document()
    _set_doc_font(doc)

    # ---- 표지/헤더
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("기능 테스트 보고서")
    run.font.size = Pt(22)
    run.font.bold = True
    _kfont(run)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"{data['filename']}  ·  Metis FinOps Test-Report Agent")
    r2.font.size = Pt(11)
    r2.font.color.rgb = MUTED

    verdict = data["verdict"]
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(f"종합 판정: {verdict}  (종합 {data['total_score']}점)")
    vr.font.size = Pt(14)
    vr.font.bold = True
    vr.font.color.rgb = GREEN if "통과" in verdict and "조건부" not in verdict else \
        (RGBColor(0xB8, 0x86, 0x0B) if "조건부" in verdict else RED)

    _h(doc, "1. 개요", 1)
    _table(doc, ["항목", "내용"], [
        ["생성 시각", data["ts"]],
        ["분석 대상", f"{data['filename']} ({data['static']['lines']}줄, {data['language'].upper()})"],
        ["분석 에이전트", "test-report-agent (Metis Gateway 경유)"],
        ["Run ID", data["run_id"]],
        ["LLM 리뷰 모드", data["mode"]],
        ["동적 테스트 방식", data["dynamic"].get("kind", "-")],
    ], widths=[1.6, 4.6])

    _h(doc, "2. 점수 요약", 1)
    doc.add_picture(_score_chart(data["scores"]), width=Inches(5.8))
    _table(doc, ["영역", "점수", "비고"], [
        ["구조 품질", data["scores"]["struct"], f"경고 {data['n_warn']}건, 정보 {data['n_info']}건"],
        ["안전성", data["scores"]["safety"], f"심각 이슈 {data['n_crit']}건"],
        ["동적 테스트", data["scores"]["test"], f"pass {data['n_pass']} / fail {data['n_fail']} / skip {data['n_skip']}"],
        ["문서화", data["scores"]["doc"], f"주석/문서화 커버리지 {data['doc_cov_pct']}"],
    ], widths=[1.6, 1.0, 3.6])

    _h(doc, "3. 정적 분석", 1)
    sa = data["static"]
    if not sa["syntax_ok"]:
        p = doc.add_paragraph()
        r = p.add_run(f"구문 오류: {sa['syntax_error']}")
        r.font.color.rgb = RED
    if sa["functions"]:
        _table(doc, ["함수/메서드", "라인", "인자", "문서화"],
               [[f["name"], f["line"], f["args"], "O" if f.get("doc") else "X"]
                for f in sa["functions"][:30]], widths=[2.6, 0.9, 0.9, 1.0])
    if sa["classes"]:
        doc.add_paragraph(f"클래스: {', '.join(c['name'] for c in sa['classes'])}")

    _h(doc, "4. 발견된 이슈", 1)
    if sa["issues"]:
        sev_ko = {"critical": "심각", "warning": "경고", "info": "정보"}
        rows = [[sev_ko[i["sev"]], i["line"] or "-", i["msg"]]
                for i in sorted(sa["issues"], key=lambda x: {"critical": 0, "warning": 1, "info": 2}[x["sev"]])]
        t = _table(doc, ["심각도", "라인", "내용"], rows, widths=[0.9, 0.8, 4.5])
        for ridx, i in enumerate(sorted(sa["issues"], key=lambda x: {"critical": 0, "warning": 1, "info": 2}[x["sev"]]), start=1):
            color = RED if i["sev"] == "critical" else (RGBColor(0xB8, 0x86, 0x0B) if i["sev"] == "warning" else MUTED)
            for p in t.rows[ridx].cells[0].paragraphs:
                for r in p.runs:
                    r.font.color.rgb = color
                    r.font.bold = True
    else:
        doc.add_paragraph("발견된 이슈 없음")

    _h(doc, "5. 동적 테스트", 1)
    dy = data["dynamic"]
    doc.add_paragraph(f"실행 방식: {dy.get('kind','-')}  ·  결과: {'성공' if dy.get('ok') else '실패'}"
                      + (f"  ·  오류: {dy.get('error')}" if dy.get("error") else ""))
    for n in dy.get("notes", []):
        p = doc.add_paragraph("※ " + n)
        for r in p.runs:
            r.font.color.rgb = MUTED
    if dy.get("doctests") and dy["doctests"]["attempted"]:
        doc.add_paragraph(f"doctest: {dy['doctests']['attempted']}건 중 {dy['doctests']['failed']}건 실패")
    if dy.get("func_tests"):
        _table(doc, ["테스트", "결과", "상세"],
               [[t_["name"], t_["status"], t_["detail"]] for t_ in dy["func_tests"]],
               widths=[1.8, 0.9, 3.5])

    _h(doc, "6. LLM 코드 리뷰", 1)
    for k, label in (("summary", "6.1 요약"), ("risk", "6.2 리스크 평가"), ("recommend", "6.3 개선 권고")):
        _h(doc, label, 2)
        doc.add_paragraph(data["reviews"][k])

    _h(doc, "7. FinOps 텔레메트리 (이 보고서 생성 비용)", 1)
    tel = data.get("telemetry")
    if tel and tel.get("detail"):
        doc.add_picture(_cost_chart(tel["detail"]), width=Inches(5.6))
        doc.add_picture(_token_chart(tel["detail"]), width=Inches(5.6))
        _table(doc, ["항목", "값"], [
            ["LLM 호출", f"{tel['steps']}회"],
            ["총 토큰", f"{tel['tokens']:,}"],
            ["총 비용", f"${tel['cost']:.6f}"],
            ["절감액", f"${tel['savings']:.6f}"],
        ], widths=[1.8, 3.0])
        _table(doc, ["스텝", "모델", "토큰(in/out)", "비용(USD)"],
               [[f"#{s['step']} {s['task_type']}", s["model"],
                 f"{s['input_tokens']}/{s['output_tokens']}", f"{s['cost_usd']:.6f}"]
                for s in tel["detail"]], widths=[1.8, 1.8, 1.4, 1.2])
    else:
        doc.add_paragraph("텔레메트리 조회 실패")

    foot = doc.add_paragraph()
    fr = foot.add_run("\nMetis FinOps Test-Report Agent — 모든 LLM 호출은 FinOps Gateway 를 경유하여 비용이 계측·통제됩니다.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED

    doc.save(out_path)
    return out_path
